from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT_DIR = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT_DIR / "reports" / "实验报告.docx"
SCREENSHOT_DIR = ROOT_DIR / "reports" / "screenshots"

SCREENSHOTS = [
    (
        SCREENSHOT_DIR / "smoke_test_output.png",
        "图 B-1 本地烟雾测试运行截图。截图显示依赖导入、数据库表数量、比赛数量和 SQL 安全检查均通过。",
    ),
    (
        SCREENSHOT_DIR / "deepseek_agent_output.png",
        "图 B-2 DeepSeek Agent 自然语言查询运行截图。截图显示系统生成只读 SQL、执行查询并返回中文分析结论。",
    ),
]


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(89, 89, 89)


def remove_existing_appendix(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() != "附录 B：运行截图":
            continue
        body = paragraph._element.getparent()
        start_index = body.index(paragraph._element)
        for element in list(body)[start_index:]:
            body.remove(element)
        return


def main() -> None:
    doc = Document(DOCX_PATH)
    remove_existing_appendix(doc)

    doc.add_page_break()
    doc.add_heading("附录 B：运行截图", level=1)
    doc.add_paragraph(
        "以下截图用于证明项目可以在本地环境实际运行。图表截图已在正文“Streamlit 界面与可视化”部分展示，"
        "本附录补充烟雾测试和 DeepSeek Agent 端到端查询链路截图。"
    )

    for image_path, caption in SCREENSHOTS:
        if not image_path.exists():
            raise FileNotFoundError(image_path)
        doc.add_picture(str(image_path), width=Inches(6.2))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, caption)

    doc.save(DOCX_PATH)
    print(f"Updated {DOCX_PATH}")


if __name__ == "__main__":
    main()
